import pytest
from pathlib import Path
import docx
from src.anonymizers.entity_store import EntityMappingStore
from src.anonymizers.fake_operators import StatefulFakerOperator
from src.recognizers.company_recognizer import CompanyNameRecognizer
from src.recognizers.ssn_recognizer import SSNRecognizer
from src.recognizers.dob_recognizer import DOBRecognizer


@pytest.fixture
def entity_store(tmp_path) -> EntityMappingStore:
    """
    Returns an in-memory EntityMappingStore instance backed by a temporary file path.
    """
    vault_file = tmp_path / "test_vault.json"
    return EntityMappingStore(persistence_path=str(vault_file))


@pytest.fixture
def stateful_faker(entity_store) -> StatefulFakerOperator:
    """
    Returns a StatefulFakerOperator initialized with the test EntityMappingStore.
    """
    return StatefulFakerOperator(entity_store=entity_store)


@pytest.fixture
def sample_docx_path(tmp_path) -> Path:
    """
    Creates a rich sample .docx document containing body paragraphs, split runs,
    multi-column tables, section headers, and section footers with all target PII types.
    """
    doc_path = tmp_path / "sample_test.docx"
    doc = docx.Document()

    # Section Header & Footer
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Confidential Report prepared for Acme Corp. | Contact: john.doe@acme.com"

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Page 1 - Employee SSN: 123-45-6789 | DOB: 05/14/1988"

    # Body Paragraph with contiguous text
    doc.add_paragraph(
        "Mr. Alexander Hamilton resides at 742 Evergreen Terrace, Springfield. "
        "His primary contact number is 555-0199 and IP address is 192.168.1.100."
    )

    # Body Paragraph with intentionally split runs (to test TextRunWalker)
    p_split = doc.add_paragraph()
    r1 = p_split.add_run("Please process the invoice from ")
    r2 = p_split.add_run("Global ")
    r2.bold = True
    r3 = p_split.add_run("Solutions ")
    r3.italic = True
    r4 = p_split.add_run("Inc.")
    r4.underline = True

    # Multi-row Table with PII
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Employee Name"
    table.cell(0, 1).text = "Credit Card Number"
    table.cell(1, 0).text = "Jane Smith"
    table.cell(1, 1).text = "4111-2222-3333-4444"

    doc.save(doc_path)
    return doc_path
