import pytest
import docx
from src.document_processors.text_walker import TextRunWalker
from src.document_processors.docx_processor import DocxProcessor
from src.recognizers.analyzer_factory import AnalyzerFactory
from src.anonymizers.anonymizer_factory import AnonymizerFactory
from src.anonymizers.entity_store import EntityMappingStore


def test_text_walker_extract_runs(tmp_path):
    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run("Hello ")
    p.add_run("world")
    p.add_run("!")
    
    full_text, run_spans = TextRunWalker.extract_runs_and_text(p)
    assert full_text == "Hello world!"
    assert len(run_spans) == 3
    assert run_spans[0][:2] == (0, 6)
    assert run_spans[1][:2] == (6, 11)
    assert run_spans[2][:2] == (11, 12)


def test_text_walker_apply_replacements_across_runs(tmp_path):
    doc = docx.Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Please contact ")
    r2 = p.add_run("Alex")
    r2.bold = True
    r3 = p.add_run("ander ")
    r3.italic = True
    r4 = p.add_run("Vance immediately.")
    
    # "Alexander Vance" spans across r2 ("Alex"), r3 ("ander "), and part of r4 ("Vance")
    # Character indices in "Please contact Alexander Vance immediately.":
    # r1 (0..15), r2 (15..19), r3 (19..25), r4 (25..43)
    # PII span "Alexander Vance" -> start=15, end=30
    full_text, _ = TextRunWalker.extract_runs_and_text(p)
    assert full_text[15:30] == "Alexander Vance"

    # Apply replacement with "Robert Smith"
    rep_count = TextRunWalker.apply_replacements(p, [(15, 30, "Robert Smith")])
    assert rep_count == 1

    # Check updated runs and formatting
    new_full_text, _ = TextRunWalker.extract_runs_and_text(p)
    assert new_full_text == "Please contact Robert Smith immediately."
    
    # Check that bolding on r2 was preserved and r3 was emptied safely
    assert p.runs[1].bold is True
    assert p.runs[1].text == "Robert Smith"
    assert p.runs[2].text == ""


def test_docx_processor_tables_headers_footers(sample_docx_path, tmp_path):
    out_path = tmp_path / "processed_test.docx"
    processor = DocxProcessor()
    processor.load(sample_docx_path)
    
    # Initialize real engines
    analyzer = AnalyzerFactory.get_engine()
    anonymizer = AnonymizerFactory.get_engine()
    store = EntityMappingStore()
    operators = AnonymizerFactory.get_operators(store)

    metrics = processor.process_and_redact(
        analyzer_engine=analyzer,
        anonymizer_engine=anonymizer,
        operators=operators,
        score_threshold=0.45,
    )
    processor.save(out_path)

    assert out_path.exists()
    assert metrics["total_replacements"] > 0
    
    # Reload saved document and verify original PII strings are gone
    reloaded_doc = docx.Document(out_path)
    # Check header
    header_text = reloaded_doc.sections[0].header.paragraphs[0].text
    assert "john.doe@acme.com" not in header_text
    
    # Check footer
    footer_text = reloaded_doc.sections[0].footer.paragraphs[0].text
    assert "123-45-6789" not in footer_text
    
    # Check table cell
    table_cell_text = reloaded_doc.tables[0].cell(1, 1).text
    assert "4111-2222-3333-4444" not in table_cell_text
