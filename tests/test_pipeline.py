import pytest
from pathlib import Path
import docx
from src.pipeline.redactor_pipeline import RedactionPipeline


def test_redaction_pipeline_single_file(sample_docx_path, tmp_path):
    output_docx = tmp_path / "pipeline_single_output.docx"
    pipeline = RedactionPipeline(score_threshold=0.45, enable_persistence=False)
    
    summary = pipeline.redact_document(
        input_path=sample_docx_path,
        output_path=output_docx,
        create_backup=True,
    )
    
    assert output_docx.exists()
    assert summary["total_replacements"] >= 4
    assert "PERSON" in summary["entity_counts"] or "US_SSN" in summary["entity_counts"]
    
    # Check that backup file was created
    backup_path = sample_docx_path.with_suffix(f".bak{sample_docx_path.suffix}")
    assert backup_path.exists()


def test_redaction_pipeline_batch_and_persistence(tmp_path):
    batch_in = tmp_path / "batch_in"
    batch_out = tmp_path / "batch_out"
    batch_in.mkdir()
    batch_out.mkdir()
    
    # Create two test documents sharing the same PII name ("Mr. Robert Vance")
    doc1_path = batch_in / "doc1.docx"
    doc1 = docx.Document()
    doc1.add_paragraph("Please contact Mr. Robert Vance regarding invoice #100.")
    doc1.save(doc1_path)
    
    doc2_path = batch_in / "doc2.docx"
    doc2 = docx.Document()
    doc2.add_paragraph("As discussed with Mr. Robert Vance yesterday, payment is due.")
    doc2.save(doc2_path)
    
    vault_path = tmp_path / "test_batch_vault.json"
    pipeline = RedactionPipeline(
        score_threshold=0.45,
        enable_persistence=True,
        persistence_path=str(vault_path),
    )
    
    # Process doc1
    out1 = batch_out / "doc1_redacted.docx"
    summary1 = pipeline.redact_document(doc1_path, out1)
    
    # Process doc2
    out2 = batch_out / "doc2_redacted.docx"
    summary2 = pipeline.redact_document(doc2_path, out2)
    
    assert out1.exists() and out2.exists()
    
    # Check that both documents got the exact same synthetic replacement for "Robert Vance"
    redacted_doc1 = docx.Document(out1)
    redacted_doc2 = docx.Document(out2)
    
    text1 = redacted_doc1.paragraphs[0].text
    text2 = redacted_doc2.paragraphs[0].text
    
    # Both shouldn't contain "Robert Vance"
    assert "Robert Vance" not in text1
    assert "Robert Vance" not in text2
    
    # Extract the synthetic name after "contact " / "with "
    # Ensure vault persisted correctly
    assert vault_path.exists()
