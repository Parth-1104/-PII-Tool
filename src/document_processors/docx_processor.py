from pathlib import Path
from typing import Union, Dict, Any, List, Optional
import docx
from src.document_processors.base_processor import BaseDocumentProcessor
from src.document_processors.text_walker import TextRunWalker
from src.utils.logger import get_logger
from src.utils.file_utils import validate_docx_path, ensure_output_directory

logger = get_logger("DocxProcessor")


class DocxProcessor(BaseDocumentProcessor):
    """
    Industrial-grade DOCX Document Processor built over python-docx.
    Traverses all document text containers: paragraphs, multi-column tables, nested tables,
    section headers, and section footers.
    
    Coordinates with Presidio Analyzer and Anonymizer to detect and replace PII
    while preserving exact paragraph styling and XML structural attributes.
    """

    def __init__(self):
        self.file_path: Optional[Path] = None
        self.doc: Optional[docx.Document] = None

    def load(self, file_path: Union[str, Path]) -> None:
        """
        Loads the .docx file into memory using python-docx.
        """
        self.file_path = validate_docx_path(file_path)
        try:
            self.doc = docx.Document(self.file_path)
            logger.info(f"Loaded document: {self.file_path.name}")
        except Exception as e:
            logger.error(f"Failed to open DOCX document at {file_path}: {e}")
            raise

    def process_and_redact(
        self,
        analyzer_engine: Any,
        anonymizer_engine: Any,
        operators: Dict[str, Any],
        score_threshold: float = 0.65,
    ) -> Dict[str, Any]:
        """
        Traverses all text structures inside the loaded DOCX document, executes PII analysis,
        and applies synthetic replacements via TextRunWalker.
        """
        if not self.doc:
            raise RuntimeError("No document loaded. Call load() before process_and_redact().")

        total_replacements = 0
        paragraphs_processed = 0

        # 1. Process document body paragraphs
        logger.debug("Scanning body paragraphs...")
        for p in self.doc.paragraphs:
            rep_count = self._process_paragraph(
                p, analyzer_engine, anonymizer_engine, operators, score_threshold
            )
            if rep_count > 0:
                total_replacements += rep_count
            paragraphs_processed += 1

        # 2. Process document body tables (including nested tables)
        logger.debug("Scanning document tables...")
        for table in self.doc.tables:
            total_replacements += self._process_table(
                table, analyzer_engine, anonymizer_engine, operators, score_threshold
            )

        # 3. Process section headers and footers
        logger.debug("Scanning headers and footers across sections...")
        for section_idx, section in enumerate(self.doc.sections):
            # Header
            if not section.header.is_linked_to_previous or section_idx == 0:
                for hp in section.header.paragraphs:
                    total_replacements += self._process_paragraph(
                        hp, analyzer_engine, anonymizer_engine, operators, score_threshold
                    )
                for h_table in section.header.tables:
                    total_replacements += self._process_table(
                        h_table, analyzer_engine, anonymizer_engine, operators, score_threshold
                    )

            # Footer
            if not section.footer.is_linked_to_previous or section_idx == 0:
                for fp in section.footer.paragraphs:
                    total_replacements += self._process_paragraph(
                        fp, analyzer_engine, anonymizer_engine, operators, score_threshold
                    )
                for f_table in section.footer.tables:
                    total_replacements += self._process_table(
                        f_table, analyzer_engine, anonymizer_engine, operators, score_threshold
                    )

        logger.info(
            f"Redaction completed across {paragraphs_processed} body paragraphs. "
            f"Total PII replacements applied: {total_replacements}"
        )
        return {"total_replacements": total_replacements, "paragraphs_processed": paragraphs_processed}

    def _process_table(
        self,
        table: Any,
        analyzer_engine: Any,
        anonymizer_engine: Any,
        operators: Dict[str, Any],
        score_threshold: float,
    ) -> int:
        """
        Recursively processes all cells and nested tables within a DOCX table.
        """
        table_replacements = 0
        for row in table.rows:
            for cell in row.cells:
                # Process paragraphs inside cell
                for p in cell.paragraphs:
                    table_replacements += self._process_paragraph(
                        p, analyzer_engine, anonymizer_engine, operators, score_threshold
                    )
                # Recursively process nested tables
                for nested_table in cell.tables:
                    table_replacements += self._process_table(
                        nested_table, analyzer_engine, anonymizer_engine, operators, score_threshold
                    )
        return table_replacements

    def _process_paragraph(
        self,
        paragraph: Any,
        analyzer_engine: Any,
        anonymizer_engine: Any,
        operators: Dict[str, Any],
        score_threshold: float,
    ) -> int:
        """
        Analyzes a single paragraph using TextRunWalker and Presidio engines,
        and applies offset-safe synthetic replacements across runs.
        """
        full_text, run_spans = TextRunWalker.extract_runs_and_text(paragraph)
        if not full_text or not full_text.strip():
            return 0

        # Run Presidio Analyzer
        analyzer_results = analyzer_engine.analyze(text=full_text, language="en")
        if not analyzer_results:
            return 0

        # Ensure every single result is converted to a proper RecognizerResult before passing forward
        from presidio_analyzer import RecognizerResult, EntityRecognizer, AnalysisExplanation
        from src.utils.stop_words import TARGET_PII_ENTITIES, should_ignore_candidate

        clean_filtered = []
        for item in analyzer_results:
            raw_type = getattr(item, "entity_type", getattr(item, "label_", "DEFAULT"))
            etype = str(raw_type).upper()
            if etype in ("ORG", "ORGANIZATION"):
                etype = "COMPANY_NAME"
            elif etype in ("GPE", "ADDRESS"):
                etype = "LOCATION"

            # 1. Strictly filter for target PII entities only
            if etype not in TARGET_PII_ENTITIES and etype != "DEFAULT":
                continue

            st = getattr(item, "start", getattr(item, "start_char", 0))
            en = getattr(item, "end", getattr(item, "end_char", 0))
            sc = getattr(item, "score", 0.65)

            if sc < score_threshold:
                continue

            span_text = full_text[int(st) : int(en)]
            # 2. Check against common English vocabulary & stop-words
            if should_ignore_candidate(etype, span_text):
                continue

            meta = getattr(item, "recognition_metadata", None)
            if not isinstance(meta, dict):
                meta = {
                    RecognizerResult.RECOGNIZER_NAME_KEY: getattr(item, "recognizer_name", "PresidioRecognizer"),
                    RecognizerResult.IS_SCORE_ENHANCED_BY_CONTEXT_KEY: False,
                }
            explanation = getattr(item, "analysis_explanation", None)
            if not explanation or not hasattr(explanation, "set_supportive_context_word"):
                explanation = AnalysisExplanation(
                    recognizer=meta.get(RecognizerResult.RECOGNIZER_NAME_KEY, "PresidioRecognizer"),
                    original_score=float(sc),
                    textual_explanation=f"Detected {etype}",
                )

            clean_filtered.append(
                RecognizerResult(
                    entity_type=etype,
                    start=int(st),
                    end=int(en),
                    score=float(sc),
                    analysis_explanation=explanation,
                    recognition_metadata=meta,
                )
            )

        if not clean_filtered:
            return 0

        # Remove overlapping spans from analyzer results (keep highest confidence/longest)
        filtered_results = EntityRecognizer.remove_duplicates(clean_filtered)

        # Run Presidio Anonymizer with custom Faker operators
        # This triggers StatefulFakerOperator for each span and updates our mapping vault
        _ = anonymizer_engine.anonymize(
            text=full_text,
            analyzer_results=filtered_results,
            operators=operators,
        )

        # Extract the exact replacement for each span from our custom operators / store
        # To get the synthetic value safely, we inspect the operator configuration callback
        replacements = []
        for res in filtered_results:
            original_substring = full_text[res.start : res.end]
            op_config = operators.get(res.entity_type)
            if not op_config and res.entity_type == "DEFAULT":
                op_config = operators.get("DEFAULT")
            
            if op_config and "lambda" in op_config.params:
                synthetic_text = op_config.params["lambda"](original_substring)
            else:
                synthetic_text = f"[{res.entity_type}]"

            replacements.append((res.start, res.end, synthetic_text))

        # Project replacements onto the paragraph runs safely
        return TextRunWalker.apply_replacements(paragraph, replacements)

    def save(self, output_path: Union[str, Path]) -> None:
        """
        Saves the modified DOCX document to the target output path.
        """
        if not self.doc:
            raise RuntimeError("No document loaded to save.")
        path = ensure_output_directory(output_path)
        try:
            self.doc.save(path)
            logger.info(f"Saved redacted document to: {path}")
        except Exception as e:
            logger.error(f"Failed to save redacted document at {path}: {e}")
            raise
